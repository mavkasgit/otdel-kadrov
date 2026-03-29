"""
HRMS - Document Generator
Генерация приказов из шаблонов
"""
import os
import sys
from datetime import datetime
from pathlib import Path

project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if project_root not in sys.path:
    sys.path.insert(0, project_root)

import settings
from core.exceptions import TemplateNotFoundError, DocumentSaveError
from core.logger import logger


class DocumentGenerator:
    """Генератор документов из шаблонов"""
    
    def __init__(self):
        self.templates_dir = settings.TEMPLATES_DIR
        self.reports_dir = settings.REPORTS_DIR
        self.templates_dir.mkdir(exist_ok=True)
        self.reports_dir.mkdir(exist_ok=True)
    
    def load_template(self, template_name: str):
        """Загрузка шаблона"""
        template_path = self.templates_dir / template_name
        if not template_path.exists():
            raise TemplateNotFoundError(f"Шаблон не найден: {template_name}")
        return template_path
    
    def populate_template(self, template_name: str, context: dict):
        """Заполнение шаблона данными"""
        try:
            from docxtpl import DocxTemplate
        except ImportError:
            logger.error("docxtpl not installed. Install: pip install docxtpl")
            raise ImportError("Для генерации документов установите docxtpl: pip install docxtpl")
        
        template_path = self.load_template(template_name)
        doc = DocxTemplate(str(template_path))
        doc.render(context)
        return doc
    
    def generate_filename(self, order_type: str, employee_name: str, order_date: datetime) -> str:
        """Генерация имени файла: {тип} {ФИО} {дата}.docx"""
        clean_name = employee_name.replace("_", " ").replace("/", " ").replace("\\", " ")
        date_str = order_date.strftime("%Y-%m-%d")
        filename = f"{order_type} {clean_name} {date_str}.docx"
        invalid_chars = '<>:"*?|'
        for char in invalid_chars:
            filename = filename.replace(char, "")
        return filename
    
    def save_document(self, doc, filename: str) -> str:
        """Сохранение документа"""
        filepath = self.reports_dir / filename
        try:
            doc.save(str(filepath))
            logger.info(f"Document saved: {filepath}")
            return str(filepath)
        except Exception as e:
            raise DocumentSaveError(f"Не удалось сохранить документ: {e}")
    
    def generate_order(self, order_type: str, employee_data: dict, order_number: str, 
                       order_date: datetime, template_name: str = None) -> str:
        """Генерация приказа"""
        if template_name is None:
            template_name = "prikaz.docx"
        
        context = {
            "order_number": order_number,
            "order_date": order_date.strftime("%d.%m.%Y"),
            "order_date_full": order_date.strftime("%d %B %Y"),
            "order_type": order_type,
            "tab_number": employee_data.get("Таб. №", ""),
            "full_name": employee_data.get("ФИО", ""),
            "position": employee_data.get("Должность", ""),
            "department": employee_data.get("Подразделение", ""),
            "hire_date": employee_data.get("Дата принятия", ""),
            "contract_start": employee_data.get("Начало контракта", ""),
            "contract_end": employee_data.get("Конец контракта", ""),
            "passport": employee_data.get("№ паспорта", ""),
            "personal_number": employee_data.get("Личный №", ""),
            "current_date": datetime.now().strftime("%d.%m.%Y"),
        }
        
        try:
            doc = self.populate_template(template_name, context)
            filename = self.generate_filename(order_type, employee_data.get("ФИО", ""), order_date)
            filepath = self.save_document(doc, filename)
            logger.info(f"Generated order: {order_number} for {employee_data.get('ФИО')}")
            return filepath
        except TemplateNotFoundError:
            logger.warning(f"Template not found: {template_name}, creating simple document")
            return self._create_simple_doc(order_type, employee_data, order_number, order_date)
    
    def _create_simple_doc(self, order_type: str, employee_data: dict, 
                           order_number: str, order_date: datetime) -> str:
        """Создание простого документа если шаблон не найден"""
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed")
            return ""
        
        doc = Document()
        doc.add_heading(f"Приказ № {order_number}", 0)
        doc.add_paragraph(f"Дата: {order_date.strftime('%d.%m.%Y')}")
        doc.add_paragraph(f"Тип: {order_type}")
        doc.add_paragraph("")
        doc.add_heading("Сотрудник:", level=1)
        
        for key, value in employee_data.items():
            if value:
                doc.add_paragraph(f"{key}: {value}")
        
        filename = self.generate_filename(order_type, employee_data.get("ФИО", ""), order_date)
        return self.save_document(doc, filename)


def main():
    """Тест генерации"""
    generator = DocumentGenerator()
    
    test_employee = {
        "Таб. №": 1,
        "ФИО": "Иванов Иван Иванович",
        "Должность": "Инженер",
        "Подразделение": "Завод КТМ",
        "Дата принятия": "01.01.2020",
    }
    
    test_date = datetime.now()
    
    try:
        filepath = generator.generate_order(
            "Прием на работу",
            test_employee,
            "001-П",
            test_date
        )
        print(f"Generated: {filepath}")
    except Exception as e:
        print(f"Error: {e}")


if __name__ == "__main__":
    main()
